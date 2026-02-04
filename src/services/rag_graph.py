import asyncio
import httpx
import faiss
import numpy as np
from typing import List, Dict, Optional, Set, Tuple
import networkx as nx
from collections import defaultdict
import tiktoken
from community import community_louvain
import json
import pickle
from pathlib import Path
from datetime import datetime
import hashlib
import uuid
import logging
from dataclasses import dataclass
from difflib import SequenceMatcher
from pypdf import PdfReader
import pdfplumber
from docx import Document as DocxDocument
import pandas as pd
import re

logger = logging.getLogger(__name__)


@dataclass
class EntityAlignment:
    """实体对齐结果"""
    canonical_name: str  # 标准名称
    aliases: List[str]   # 别名列表
    similarity: float    # 相似度


class AsyncLLMClient:
    """异步 LLM 客户端"""
    
    def __init__(self, base_url: str, api_key: str, model: str, max_concurrent: int = 10):
        self.base_url = base_url.rstrip('/')
        self.api_key = api_key
        self.model = model
        self.semaphore = asyncio.Semaphore(max_concurrent)
        
        # 使用连接池
        self.client = httpx.AsyncClient(
            timeout=httpx.Timeout(60.0, connect=10.0),
            limits=httpx.Limits(max_keepalive_connections=20, max_connections=100)
        )
    
    async def chat(self, messages: List[Dict], temperature: float = 0, 
                   max_tokens: int = 10000, response_format: Optional[Dict] = None) -> str:
        """异步聊天补全"""
        async with self.semaphore:
            payload = {
                "model": self.model,
                "messages": messages,
                "temperature": temperature,
                "max_tokens": max_tokens
            }
            
            if response_format:
                payload["response_format"] = response_format
            
            try:
                response = await self.client.post(
                    f"{self.base_url}/chat/completions",
                    headers={
                        "Authorization": f"Bearer {self.api_key}",
                        "Content-Type": "application/json"
                    },
                    json=payload
                )
                response.raise_for_status()
                result = response.json()
                return result['choices'][0]['message']['content']
            except Exception as e:
                logger.error(f"LLM 调用失败: {e}")
                raise
    
    async def close(self):
        """关闭客户端"""
        await self.client.aclose()


class AsyncEmbeddingClient:
    """异步 Embedding 客户端"""
    
    def __init__(self, base_url: str, api_key: str, model: str, max_concurrent: int = 20):
        self.base_url = base_url.rstrip('/')
        self.api_key = api_key
        self.model = model
        self.semaphore = asyncio.Semaphore(max_concurrent)
        
        self.client = httpx.AsyncClient(
            timeout=httpx.Timeout(60.0, connect=10.0),
            limits=httpx.Limits(max_keepalive_connections=20, max_connections=100)
        )
    
    async def embed(self, texts: List[str]) -> List[np.ndarray]:
        """批量生成 embeddings"""
        async with self.semaphore:
            try:
                response = await self.client.post(
                    f"{self.base_url}/embeddings",
                    headers={
                        "Authorization": f"Bearer {self.api_key}",
                        "Content-Type": "application/json"
                    },
                    json={
                        "model": self.model,
                        "input": texts
                    }
                )
                response.raise_for_status()
                result = response.json()
                return [np.array(item['embedding'], dtype='float32') 
                       for item in result['data']]
            except Exception as e:
                logger.error(f"Embedding 调用失败: {e}")
                raise
    
    async def close(self):
        """关闭客户端"""
        await self.client.aclose()


class DocumentParser:
    """文档解析器 (保持同步，IO 密集型)"""
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """解析 PDF 文件"""
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page_num, page in enumerate(pdf.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(f"[Page {page_num}]\n{page_text}")
                    
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, 1):
                        if table:
                            table_text = f"\n[Table {table_num} on Page {page_num}]\n"
                            for row in table:
                                table_text += " | ".join([str(cell) if cell else "" for cell in row]) + "\n"
                            text_parts.append(table_text)
                
                return "\n\n".join(text_parts)
        except Exception as e:
            logger.warning(f"pdfplumber 解析失败，使用 pypdf: {e}")
            reader = PdfReader(file_path)
            text_parts = []
            for page_num, page in enumerate(reader.pages, 1):
                text = page.extract_text()
                if text:
                    text_parts.append(f"[Page {page_num}]\n{text}")
            return "\n\n".join(text_parts)
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """解析 DOCX 文件"""
        doc = DocxDocument(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        for table_num, table in enumerate(doc.tables, 1):
            table_text = f"\n[Table {table_num}]\n"
            for row in table.rows:
                row_text = " | ".join([cell.text for cell in row.cells])
                table_text += row_text + "\n"
            text_parts.append(table_text)
        
        return "\n\n".join(text_parts)
    
    @staticmethod
    def parse_txt(file_path: str) -> str:
        """解析纯文本文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    @staticmethod
    def parse_markdown(file_path: str) -> str:
        """解析 Markdown 文件"""
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    @staticmethod
    def parse_csv(file_path: str) -> str:
        """解析 CSV 文件"""
        df = pd.read_csv(file_path)
        text = f"CSV Data ({len(df)} rows x {len(df.columns)} columns)\n\n"
        text += df.to_string(index=False)
        return text
    
    @classmethod
    def parse_document(cls, file_path: str) -> str:
        """根据文件扩展名自动选择解析器"""
        path = Path(file_path)
        extension = path.suffix.lower()
        
        parsers = {
            '.pdf': cls.parse_pdf,
            '.docx': cls.parse_docx,
            '.doc': cls.parse_docx,
            '.txt': cls.parse_txt,
            '.md': cls.parse_markdown,
            '.markdown': cls.parse_markdown,
            '.csv': cls.parse_csv,
        }
        
        parser = parsers.get(extension)
        if parser is None:
            raise ValueError(f"不支持的文件格式: {extension}")
        
        return parser(str(path))


class TextChunker:
    """智能文本分块器"""
    
    def __init__(self, chunk_size: int = 600, overlap: int = 100):
        self.chunk_size = chunk_size
        self.overlap = overlap
        self.encoding = tiktoken.get_encoding("cl100k_base")
    
    def chunk_by_sentences(self, text: str) -> List[str]:
        """按句子分块（保持句子完整性）"""
        sentences = re.split(r'[.!?。！？]\s+', text)
        
        chunks = []
        current_chunk = []
        current_length = 0
        
        for sentence in sentences:
            sentence_words = len(sentence.split())
            
            if current_length + sentence_words > self.chunk_size and current_chunk:
                chunks.append(' '.join(current_chunk))
                
                overlap_sentences = []
                overlap_length = 0
                for s in reversed(current_chunk):
                    s_length = len(s.split())
                    if overlap_length + s_length <= self.overlap:
                        overlap_sentences.insert(0, s)
                        overlap_length += s_length
                    else:
                        break
                
                current_chunk = overlap_sentences
                current_length = overlap_length
            
            current_chunk.append(sentence)
            current_length += sentence_words
        
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def chunk_text(self, text: str) -> List[str]:
        """分块文本"""
        return self.chunk_by_sentences(text)


class EntityAligner:
    """实体对齐器 - 合并相似实体"""
    
    def __init__(self, similarity_threshold: float = 0.85):
        self.similarity_threshold = similarity_threshold
    
    def calculate_similarity(self, name1: str, name2: str) -> float:
        """计算两个实体名的相似度"""
        # 1. 完全匹配
        if name1.lower() == name2.lower():
            return 1.0
        
        # 2. 字符串相似度
        seq_similarity = SequenceMatcher(None, name1.lower(), name2.lower()).ratio()
        
        # 3. 包含关系
        if name1.lower() in name2.lower() or name2.lower() in name1.lower():
            return max(seq_similarity, 0.9)
        
        # 4. 词集相似度 (Jaccard)
        words1 = set(name1.lower().split())
        words2 = set(name2.lower().split())
        if words1 and words2:
            jaccard = len(words1 & words2) / len(words1 | words2)
            return max(seq_similarity, jaccard)
        
        return seq_similarity
    
    def align_entities(self, entities: Dict[str, Dict]) -> Dict[str, EntityAlignment]:
        """对齐实体，返回映射关系"""
        entity_names = list(entities.keys())
        alignments = {}
        processed = set()
        
        for i, name1 in enumerate(entity_names):
            if name1 in processed:
                continue
            
            # 查找相似实体
            similar_entities = [name1]
            
            for name2 in entity_names[i+1:]:
                if name2 in processed:
                    continue
                
                similarity = self.calculate_similarity(name1, name2)
                
                if similarity >= self.similarity_threshold:
                    similar_entities.append(name2)
                    processed.add(name2)
            
            # 选择最具代表性的名称作为标准名
            canonical_name = max(similar_entities, key=len)  # 选最长的
            
            alignment = EntityAlignment(
                canonical_name=canonical_name,
                aliases=similar_entities,
                similarity=1.0
            )
            
            for entity_name in similar_entities:
                alignments[entity_name] = alignment
            
            processed.add(name1)
        
        return alignments


class GraphRAGPipeline:
    """
    异步 GraphRAG Pipeline
    
    优化:
    1. 全面异步化
    2. 增量构建：新文档只触发局部更新
    3. 实体对齐：智能合并相似实体
    4. 批量并发处理
    """

    def __init__(self, llm_api_key: str, embedding_api_key: str, llm_url: str, 
                 embedding_url: str, embedding_name: str, embedding_dim: int,
                 llm_name: str, storage_dir: str = "./graphrag_storage",
                 max_llm_concurrent: int = 10, max_embed_concurrent: int = 20):
        
        self.llm_client = AsyncLLMClient(llm_url, llm_api_key, llm_name, max_llm_concurrent)
        self.embedding_client = AsyncEmbeddingClient(embedding_url, embedding_api_key, 
                                                     embedding_name, max_embed_concurrent)
        
        self.embedding_name = embedding_name
        self.llm_name = llm_name
        self.dimension = embedding_dim
        
        # 存储目录
        self.storage_dir = Path(storage_dir)
        self.storage_dir.mkdir(parents=True, exist_ok=True)
        
        # 文档管理
        self.document_parser = DocumentParser()
        self.text_chunker = TextChunker()
        self.entity_aligner = EntityAligner(similarity_threshold=0.85)
        
        self.documents = {}
        
        # UUID 映射
        self.uuid_to_docid = {}
        self.docid_to_uuid = {}
        
        # 图谱数据
        self.text_chunks = []
        self.chunk_to_doc = {}
        self.entities = {}
        self.entity_alignments = {}  # 实体对齐映射
        self.relationships = []
        self.claims = []
        
        # 知识图谱
        self.graph = nx.Graph()
        
        # 社区结构
        self.communities = {}
        self.community_summaries = {}
        
        # FAISS 索引
        self.community_summary_index = None
        self.community_embeddings = []
        
        self.encoding = tiktoken.get_encoding("cl100k_base")
        
        # 增量构建标记
        self.needs_rebuild = {
            'entities': False,
            'graph': False,
            'communities': False,
            'summaries': False,
            'index': False
        }
    
    async def initialize(self):
        """异步初始化 - 自动加载已有数据"""
        try:
            await self.load("default")
            logger.info(f"✅ 自动加载知识库成功: {len(self.documents)} 个文档, {len(self.text_chunks)} 个chunks")
        except FileNotFoundError:
            logger.info("📝 未找到已有知识库，将创建新的")
        except Exception as e:
            logger.warning(f"⚠️ 加载知识库失败: {e}")
    
    # ==================== 文档管理 ====================
    
    def _calculate_file_hash(self, file_path: str) -> str:
        """计算文件哈希"""
        with open(file_path, 'rb') as f:
            return hashlib.md5(f.read()).hexdigest()
    
    async def add_document(self, file_path: str, metadata: Optional[Dict] = None, 
                          doc_uuid: Optional[str] = None) -> str:
        """
        异步添加文档 (增量模式)
        
        Returns:
            文档的 UUID
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        if doc_uuid is None:
            doc_uuid = str(uuid.uuid4())
        
        file_hash = self._calculate_file_hash(str(file_path))
        
        # 检查文档是否已存在
        if file_hash in self.documents:
            logger.info(f"⚠️ 文档已存在: {file_path.name}")
            return self.docid_to_uuid.get(file_hash, doc_uuid)
        
        logger.info(f"📄 添加文档: {file_path.name} (UUID: {doc_uuid})")
        
        # 1. 解析文档 (在线程池中执行，避免阻塞事件循环)
        loop = asyncio.get_event_loop()
        text = await loop.run_in_executor(None, self.document_parser.parse_document, str(file_path))
        logger.info(f"  📝 文档解析完成，文本长度: {len(text)} 字符")
        
        # 2. 分块
        chunks = await loop.run_in_executor(None, self.text_chunker.chunk_text, text)
        logger.info(f"  ✂️ 分块完成: {len(chunks)} 个chunks")
        
        # 3. 记录文档信息
        doc_info = {
            'uuid': doc_uuid,
            'path': str(file_path),
            'name': file_path.name,
            'hash': file_hash,
            'chunks': chunks,
            'chunk_ids': [],
            'metadata': metadata or {},
            'added_at': datetime.now().isoformat()
        }
        
        self.documents[file_hash] = doc_info
        self.uuid_to_docid[doc_uuid] = file_hash
        self.docid_to_uuid[file_hash] = doc_uuid
        
        # 4. 异步并发提取图元素
        chunk_start_id = len(self.text_chunks)
        logger.info(f"  🔍 开始异步提取图元素 (起始ID: {chunk_start_id})...")
        
        # 并发提取所有 chunks
        tasks = []
        for chunk_id, chunk in enumerate(chunks):
            global_chunk_id = chunk_start_id + chunk_id
            task = self.extract_graph_elements(chunk, global_chunk_id)
            tasks.append(task)
        
        # 等待所有提取任务完成
        chunk_elements = await asyncio.gather(*tasks)
        
        # 5. 添加到数据结构
        for chunk_id, elements in enumerate(chunk_elements):
            global_chunk_id = chunk_start_id + chunk_id
            self.text_chunks.append(elements)
            self.chunk_to_doc[global_chunk_id] = file_hash
            doc_info['chunk_ids'].append(global_chunk_id)
        
        logger.info(f"  ✅ 完成: 提取了 {len(chunks)} 个文本块")
        logger.info(f"  📊 当前总计: {len(self.text_chunks)} 个chunks")
        
        # 6. 标记需要增量更新
        self._mark_needs_rebuild(['entities', 'graph', 'communities', 'summaries', 'index'])
        
        # 7. 自动保存
        await self.save("default")
        logger.info(f"  💾 知识库已自动保存")
        
        return doc_uuid
    
    async def remove_document(self, doc_id: str):
        """异步删除文档"""
        if doc_id in self.uuid_to_docid:
            internal_doc_id = self.uuid_to_docid[doc_id]
            doc_uuid = doc_id
        elif doc_id in self.documents:
            internal_doc_id = doc_id
            doc_uuid = self.docid_to_uuid.get(doc_id)
        else:
            raise ValueError(f"文档不存在: {doc_id}")
        
        logger.info(f"🗑️ 删除文档: {self.documents[internal_doc_id]['name']}")
        
        # 标记删除的 chunks
        chunk_ids = set(self.documents[internal_doc_id]['chunk_ids'])
        for chunk_id in chunk_ids:
            if chunk_id < len(self.text_chunks):
                self.text_chunks[chunk_id] = {'entities': [], 'relationships': [], 'claims': []}
            self.chunk_to_doc.pop(chunk_id, None)
        
        # 删除映射
        if doc_uuid:
            self.uuid_to_docid.pop(doc_uuid, None)
            self.docid_to_uuid.pop(internal_doc_id, None)
        
        del self.documents[internal_doc_id]
        
        # 标记需要重建
        self._mark_needs_rebuild(['entities', 'graph', 'communities', 'summaries', 'index'])
        
        # 自动保存
        await self.save("default")
        logger.info("  💾 删除后已自动保存")
        logger.info("  ✅ 文档已删除")
    
    def list_documents(self) -> List[Dict]:
        """列出所有文档"""
        return [
            {
                'uuid': info['uuid'],
                'id': doc_id,
                'name': info['name'],
                'path': info['path'],
                'chunks': len(info['chunks']),
                'added_at': info['added_at'],
                'metadata': info['metadata']
            }
            for doc_id, info in self.documents.items()
        ]
    
    # ==================== 图元素提取 (异步) ====================
    @staticmethod
    def safe_json_loads(text: str):
        logger.info(f"大模型结果：{text}")
        try:
            return json.loads(text)
        except json.JSONDecodeError:
            # 1️⃣ 去掉 markdown fence
            cleaned = re.sub(r"```json|```", "", text, flags=re.I).strip()
            try:
                return json.loads(cleaned)
            except json.JSONDecodeError:
                # 2️⃣ 尝试截取第一个 { ... }
                match = re.search(r"\{.*\}", cleaned, re.S)
                if match:
                    return json.loads(match.group())
                raise

    
    async def extract_graph_elements(self, text: str, chunk_id: int) -> Dict:
        """异步提取图元素"""
        
        prompt = f"""从以下文本中提取结构化信息，必须返回JSON格式，格式完整。

文本:
{text}

提取内容:
1. entities: [{{"name": "实体名", "type": "类型", "description": "描述"}}]
2. relationships: [{{"source": "源实体", "target": "目标实体", "description": "关系", "strength": 1-10}}]
3. claims: [{{"subject": "主体", "object": "客体", "type": "FACT/OPINION", "description": "描述", "date": "时间"}}]

只返回JSON，不要其他内容，格式必须完整。
"""

        try:
            response_text = await self.llm_client.chat(
                messages=[
                    {"role": "system", "content": "你是知识图谱专家。"},
                    {"role": "user", "content": prompt}
                ],
                temperature=0,
                response_format={"type": "json_object"}
            )
            
            result = self.safe_json_loads(response_text) 
            logger.debug(f"Chunk {chunk_id} 提取结果: {len(result.get('entities', []))} 实体")
            return result
            
        except Exception as e:
            logger.error(f"提取失败 (chunk {chunk_id}): {e}")
            return {"entities": [], "relationships": [], "claims": []}
    
    async def summarize_entity(self, entity_name: str, descriptions: List[str]) -> str:
        """异步合并实体描述"""
        if len(descriptions) == 1:
            return descriptions[0]
        
        combined = "\n".join([f"- {desc}" for desc in descriptions])
        
        prompt = f"""整合以下关于"{entity_name}"的描述为一个摘要（150-200词）：

{combined}

只返回摘要。"""

        response_text = await self.llm_client.chat(
            messages=[{"role": "user", "content": prompt}],
            temperature=0,
            max_tokens=300
        )
        
        return response_text.strip()
    
    # ==================== 图谱构建 (异步 + 增量) ====================
    
    def _mark_needs_rebuild(self, stages: List[str]):
        """标记需要重建的阶段"""
        for stage in stages:
            self.needs_rebuild[stage] = True
    
    async def merge_entities_and_relationships(self, incremental: bool = True):
        """
        异步合并实体和关系
        
        Args:
            incremental: 是否增量模式（仅处理新数据）
        """
        logger.info(f"📊 开始{'增量' if incremental else '全量'}合并实体和关系...")
        
        if not self.text_chunks:
            logger.warning("⚠️ text_chunks 为空！请先添加文档。")
            return
        
        # 如果是增量模式且已有实体，保留原有数据
        if incremental and self.entities:
            logger.info("  使用增量模式，保留已有实体")
            entity_descriptions = defaultdict(list, {
                name: [data['description']] for name, data in self.entities.items()
            })
            entity_types = {name: data['type'] for name, data in self.entities.items()}
            entity_sources = defaultdict(set, {
                name: set(data['source_ids']) for name, data in self.entities.items()
            })
        else:
            entity_descriptions = defaultdict(list)
            entity_types = {}
            entity_sources = defaultdict(set)
        
        # 收集新实体
        for chunk_id, chunk_data in enumerate(self.text_chunks):
            entities = chunk_data.get('entities', [])
            
            for entity in entities:
                name = entity['name']
                entity_descriptions[name].append(entity['description'])
                entity_types[name] = entity['type']
                entity_sources[name].add(chunk_id)
        
        logger.info(f"  发现 {len(entity_descriptions)} 个唯一实体")
        
        # ★★★ 实体对齐 ★★★
        logger.info("  🔄 执行实体对齐...")
        self.entity_alignments = self.entity_aligner.align_entities(
            {name: {} for name in entity_descriptions.keys()}
        )
        
        # 统计对齐结果
        aligned_groups = defaultdict(list)
        for original_name, alignment in self.entity_alignments.items():
            aligned_groups[alignment.canonical_name].append(original_name)
        
        merged_count = sum(1 for aliases in aligned_groups.values() if len(aliases) > 1)
        logger.info(f"  对齐完成: {len(entity_descriptions)} 个实体 → {len(aligned_groups)} 个标准实体 (合并了 {merged_count} 组)")
        
        # 使用对齐后的实体
        aligned_descriptions = defaultdict(list)
        aligned_types = {}
        aligned_sources = defaultdict(set)
        
        for original_name, alignment in self.entity_alignments.items():
            canonical = alignment.canonical_name
            aligned_descriptions[canonical].extend(entity_descriptions[original_name])
            aligned_types[canonical] = entity_types[original_name]
            aligned_sources[canonical].update(entity_sources[original_name])
        
        # 异步生成实体摘要
        logger.info("  生成实体摘要...")
        tasks = []
        entity_names = []
        
        for entity_name, descriptions in aligned_descriptions.items():
            if len(descriptions) > 1 or not incremental or entity_name not in self.entities:
                tasks.append(self.summarize_entity(entity_name, descriptions))
                entity_names.append(entity_name)
        
        if tasks:
            summaries = await asyncio.gather(*tasks)
            
            for entity_name, summary in zip(entity_names, summaries):
                self.entities[entity_name] = {
                    'description': summary,
                    'type': aligned_types[entity_name],
                    'source_ids': list(aligned_sources[entity_name]),
                    'aliases': [alias for alias, align in self.entity_alignments.items() 
                               if align.canonical_name == entity_name and alias != entity_name]
                }
        
        # 合并关系 (使用对齐后的实体名)
        relationship_map = defaultdict(lambda: {'descriptions': [], 'strengths': [], 'sources': set()})
        
        for chunk_id, chunk_data in enumerate(self.text_chunks):
            for rel in chunk_data.get('relationships', []):
                # 获取对齐后的实体名
                source = self.entity_alignments.get(rel['source'], 
                        EntityAlignment(rel['source'], [rel['source']], 1.0)).canonical_name
                target = self.entity_alignments.get(rel['target'], 
                        EntityAlignment(rel['target'], [rel['target']], 1.0)).canonical_name
                
                key = (source, target)
                relationship_map[key]['descriptions'].append(rel['description'])
                relationship_map[key]['strengths'].append(rel.get('strength', 5))
                relationship_map[key]['sources'].add(chunk_id)
        
        self.relationships = []
        for (source, target), data in relationship_map.items():
            if source in self.entities and target in self.entities:
                self.relationships.append({
                    'source': source,
                    'target': target,
                    'description': '; '.join(data['descriptions']),
                    'weight': float(np.mean(data['strengths'])),
                    'source_ids': list(data['sources'])
                })
        
        logger.info(f"  ✅ 完成: {len(self.entities)} 实体, {len(self.relationships)} 关系")
        self.needs_rebuild['entities'] = False
    
    async def build_graph(self):
        """异步构建知识图谱"""
        logger.info("🕸️ 构建知识图谱...")
        
        self.graph = nx.Graph()
        
        for entity_name, entity_data in self.entities.items():
            self.graph.add_node(
                entity_name,
                type=entity_data['type'],
                description=entity_data['description']
            )
        
        for rel in self.relationships:
            self.graph.add_edge(
                rel['source'],
                rel['target'],
                weight=rel['weight'],
                description=rel['description']
            )
        
        logger.info(f"  ✅ 图谱: {self.graph.number_of_nodes()} 节点, {self.graph.number_of_edges()} 边")
        self.needs_rebuild['graph'] = False
    
    def detect_hierarchical_communities(self, max_level: int = 3):
        """层次化社区检测 (保持同步，因为 community_louvain 是同步的)"""
        logger.info("👥 社区检测...")
        
        self.communities = {}
        current_graph = self.graph.copy()
        
        for level in range(max_level):
            partition = community_louvain.best_partition(
                current_graph,
                weight='weight',
                resolution=1.0
            )
            
            communities_at_level = defaultdict(list)
            for node, comm_id in partition.items():
                communities_at_level[comm_id].append(node)
            
            self.communities[level] = dict(communities_at_level)
            logger.info(f"  Level {level}: {len(communities_at_level)} 个社区")
            
            if len(communities_at_level) <= 1:
                break
            
            # 构建下一层
            next_graph = nx.Graph()
            for comm_id in communities_at_level.keys():
                next_graph.add_node(f"comm_{level}_{comm_id}")
            
            for u, v, data in current_graph.edges(data=True):
                comm_u = partition[u]
                comm_v = partition[v]
                if comm_u != comm_v:
                    edge_key = (f"comm_{level}_{comm_u}", f"comm_{level}_{comm_v}")
                    if next_graph.has_edge(*edge_key):
                        next_graph[edge_key[0]][edge_key[1]]['weight'] += data.get('weight', 1)
                    else:
                        next_graph.add_edge(*edge_key, weight=data.get('weight', 1))
            
            current_graph = next_graph
        
        self.needs_rebuild['communities'] = False
    
    async def generate_community_summary(self, level: int, community_id: int) -> str:
        """异步生成社区摘要"""
        nodes = self.communities[level][community_id]
        
        entities_info = []
        for node in nodes[:20]:
            if node in self.entities:
                entities_info.append(
                    f"- {node} ({self.entities[node]['type']}): "
                    f"{self.entities[node]['description'][:200]}"
                )
        
        relationships_info = []
        for rel in self.relationships:
            if rel['source'] in nodes and rel['target'] in nodes:
                relationships_info.append(
                    f"- {rel['source']} → {rel['target']}: {rel['description'][:150]}"
                )
        
        prompt = f"""生成社区摘要（300-400词）：

实体:
{chr(10).join(entities_info)}

关系:
{chr(10).join(relationships_info[:15])}

包括：主题、关键实体、关键发现、连接性。只返回摘要。"""

        response_text = await self.llm_client.chat(
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=500
        )
        
        return response_text.strip()
    
    async def generate_all_community_summaries(self):
        """异步生成所有社区摘要"""
        logger.info("📝 生成社区摘要...")
        
        self.community_summaries = {}
        tasks = []
        keys = []
        
        for level, communities in self.communities.items():
            for comm_id in communities.keys():
                tasks.append(self.generate_community_summary(level, comm_id))
                keys.append((level, comm_id))
        
        summaries = await asyncio.gather(*tasks)
        
        for key, summary in zip(keys, summaries):
            self.community_summaries[key] = summary
        
        logger.info(f"  ✅ 生成了 {len(self.community_summaries)} 个社区摘要")
        self.needs_rebuild['summaries'] = False
    
    async def build_community_summary_index(self):
        """异步构建向量索引"""
        logger.info("🔍 构建向量索引...")
        
        summaries = []
        summary_metadata = []
        
        for (level, comm_id), summary in self.community_summaries.items():
            summaries.append(summary)
            summary_metadata.append({
                'level': level,
                'community_id': comm_id,
                'summary': summary
            })
        
        if not summaries:
            logger.warning("⚠️ 没有社区摘要可索引")
            return
        
        # 批量生成 embeddings
        batch_size = 100
        embeddings = []
        
        logger.info(f"  生成 {len(summaries)} 个摘要的向量...")
        for i in range(0, len(summaries), batch_size):
            batch = summaries[i:i + batch_size]
            batch_embeddings = await self.embedding_client.embed(batch)
            embeddings.extend(batch_embeddings)
        
        self.community_embeddings = summary_metadata
        
        # 构建 FAISS
        embeddings_array = np.array(embeddings, dtype='float32')
        self.community_summary_index = faiss.IndexFlatIP(self.dimension)
        faiss.normalize_L2(embeddings_array)
        self.community_summary_index.add(embeddings_array)
        
        logger.info(f"  ✅ 索引完成: {len(embeddings)} 个社区")
        self.needs_rebuild['index'] = False
    
    # ==================== 索引构建 (智能增量) ====================
    
    async def rebuild_index(self, force_full: bool = False):
        """
        智能重建索引
        
        Args:
            force_full: 强制全量重建
        """
        logger.info("=" * 60)
        logger.info(f"🔄 {'全量' if force_full else '增量'}重建 GraphRAG 索引")
        logger.info("=" * 60)
        
        if not self.text_chunks:
            logger.error("❌ text_chunks 为空！请先添加文档。")
            raise RuntimeError("没有文档可以索引，请先上传文档")
        
        logger.info(f"📊 数据统计: {len(self.text_chunks)} chunks, {len(self.documents)} 文档")
        
        # 根据标记决定执行哪些阶段
        if force_full or self.needs_rebuild['entities']:
            logger.info("[1/5] 合并实体和关系...")
            await self.merge_entities_and_relationships(incremental=not force_full)
            logger.info(f"  完成: {len(self.entities)} 实体, {len(self.relationships)} 关系")
        else:
            logger.info("[1/5] 跳过 (实体已是最新)")
        
        if not self.entities:
            logger.error("❌ 没有提取到实体！请检查文档内容或 LLM 配置")
            raise RuntimeError("未能提取实体，索引构建失败")
        
        if force_full or self.needs_rebuild['graph']:
            logger.info("[2/5] 构建知识图谱...")
            await self.build_graph()
            logger.info(f"  完成: {self.graph.number_of_nodes()} 节点, {self.graph.number_of_edges()} 边")
        else:
            logger.info("[2/5] 跳过 (图谱已是最新)")
        
        if force_full or self.needs_rebuild['communities']:
            logger.info("[3/5] 社区检测...")
            loop = asyncio.get_event_loop()
            await loop.run_in_executor(None, self.detect_hierarchical_communities)
        else:
            logger.info("[3/5] 跳过 (社区已是最新)")
        
        if force_full or self.needs_rebuild['summaries']:
            logger.info("[4/5] 生成社区摘要...")
            await self.generate_all_community_summaries()
        else:
            logger.info("[4/5] 跳过 (摘要已是最新)")
        
        if force_full or self.needs_rebuild['index']:
            logger.info("[5/5] 构建向量索引...")
            await self.build_community_summary_index()
        else:
            logger.info("[5/5] 跳过 (索引已是最新)")
        
        # 自动保存
        await self.save("default")
        logger.info("💾 索引重建后已自动保存")
        
        logger.info("=" * 60)
        logger.info("✅ 索引重建完成!")
        logger.info("=" * 60)
    
    # ==================== 查询 (异步) ====================
    
    async def global_query(self, question: str, top_k_communities: int = 10, 
                          simple_mode: bool = False) -> str:
        """异步查询知识库"""
        if self.community_summary_index is None:
            raise RuntimeError("索引未构建，请先上传文档并重建索引")
        
        # 检索社区
        query_embeddings = await self.embedding_client.embed([question])
        query_embedding = np.array(query_embeddings, dtype='float32')
        faiss.normalize_L2(query_embedding)
        
        scores, indices = self.community_summary_index.search(
            query_embedding, 
            min(top_k_communities, len(self.community_embeddings))
        )
        
        # 简单模式：直接返回社区摘要
        if simple_mode:
            search_results = []
            threshold = 0.5
            
            for idx, score in zip(indices[0], scores[0]):
                if score >= threshold:
                    search_results.append(self.community_embeddings[idx]['summary'])
            
            if not search_results:
                return "抱歉，未找到相关信息。"
            
            return "\n\n".join([f"社区摘要 {i+1}\n{res}" 
                              for i, res in enumerate(search_results)])
        
        # Map-Reduce 模式
        tasks = []
        valid_indices = []
        
        for idx, score in zip(indices[0], scores[0]):
            if idx < len(self.community_embeddings):
                comm_data = self.community_embeddings[idx]
                tasks.append(self._ask_community(question, comm_data['summary']))
                valid_indices.append((idx, score, comm_data))
        
        answers = await asyncio.gather(*tasks)
        
        community_answers = []
        for (idx, score, comm_data), answer in zip(valid_indices, answers):
            if answer and len(answer.strip()) > 10:
                community_answers.append({
                    'level': comm_data['level'],
                    'community_id': comm_data['community_id'],
                    'content': answer,
                    'score': float(score)
                })
        
        return await self._reduce_answers(question, community_answers)
    
    async def _ask_community(self, question: str, community_summary: str) -> str:
        """异步询问单个社区"""
        prompt = f"""基于社区信息回答问题（2-3句话）。如果无关，回答"无相关信息"。

社区信息:
{community_summary}

问题: {question}

只返回答案。"""
        
        try:
            response_text = await self.llm_client.chat(
                messages=[{"role": "user", "content": prompt}],
                temperature=0.5,
                max_tokens=150
            )
            return response_text.strip()
        except Exception as e:
            logger.error(f"社区查询失败: {e}")
            return ""
    
    async def _reduce_answers(self, question: str, community_answers: List[Dict]) -> str:
        """异步综合答案"""
        if not community_answers:
            return "抱歉，在知识图谱中没有找到相关信息。"
        
        community_answers.sort(key=lambda x: x['score'], reverse=True)
        
        answers_text = []
        for i, ans_data in enumerate(community_answers[:10], 1):
            if ans_data['content'].lower() != "无相关信息":
                answers_text.append(f"{i}. {ans_data['content']}")
        
        if not answers_text:
            return "抱歉，找到的信息与问题不太相关。"
        
        combined = "\n".join(answers_text)
        
        prompt = f"""综合以下答案为一个连贯的最终答案（200-400词）：

问题: {question}

各社区答案:
{combined}

要求: 整合信息、消除冗余、保持清晰、呈现不同观点（如有）。

只返回最终答案。"""

        response_text = await self.llm_client.chat(
            messages=[{"role": "user", "content": prompt}],
            temperature=0.3,
            max_tokens=600
        )
        
        return response_text.strip()
    
    # ==================== 持久化 (异步) ====================
    
    async def save(self, name: str = "default"):
        """异步保存知识库"""
        save_dir = self.storage_dir / name
        save_dir.mkdir(parents=True, exist_ok=True)
        
        logger.info(f"💾 保存知识库: {save_dir}")
        
        loop = asyncio.get_event_loop()
        
        # 在线程池中执行 IO 操作
        await loop.run_in_executor(None, self._save_sync, save_dir)
        
        logger.info(f"  ✅ 保存完成: {len(self.documents)} 文档, {len(self.text_chunks)} chunks")
    
    def _save_sync(self, save_dir: Path):
        """同步保存逻辑"""
        # 保存文档
        with open(save_dir / "documents.json", 'w', encoding='utf-8') as f:
            json.dump(self.documents, f, ensure_ascii=False, indent=2)
        
        # 保存 UUID 映射
        with open(save_dir / "uuid_mappings.json", 'w', encoding='utf-8') as f:
            json.dump({
                'uuid_to_docid': self.uuid_to_docid,
                'docid_to_uuid': self.docid_to_uuid
            }, f, ensure_ascii=False, indent=2)
        
        # 保存图数据
        with open(save_dir / "graph_data.pkl", 'wb') as f:
            pickle.dump({
                'text_chunks': self.text_chunks,
                'chunk_to_doc': self.chunk_to_doc,
                'entities': self.entities,
                'entity_alignments': {k: (v.canonical_name, v.aliases, v.similarity) 
                                     for k, v in self.entity_alignments.items()},
                'relationships': self.relationships,
                'claims': self.claims,
                'communities': self.communities,
                'community_summaries': self.community_summaries,
                'community_embeddings': self.community_embeddings,
                'needs_rebuild': self.needs_rebuild,
            }, f)
        
        # 保存图
        with open(save_dir / "graph.gpickle", 'wb') as f:
            pickle.dump(self.graph, f)
        
        # 保存 FAISS
        if self.community_summary_index:
            faiss.write_index(self.community_summary_index, 
                            str(save_dir / "faiss_index.bin"))
    
    async def load(self, name: str = "default"):
        """异步加载知识库"""
        load_dir = self.storage_dir / name
        
        if not load_dir.exists():
            raise FileNotFoundError(f"知识库不存在: {load_dir}")
        
        logger.info(f"📂 加载知识库: {load_dir}")
        
        loop = asyncio.get_event_loop()
        await loop.run_in_executor(None, self._load_sync, load_dir)
        
        logger.info(f"  ✅ 加载完成: {len(self.documents)} 文档, "
                   f"{len(self.text_chunks)} chunks, {len(self.entities)} 实体")
    
    def _load_sync(self, load_dir: Path):
        """同步加载逻辑"""
        # 加载文档
        with open(load_dir / "documents.json", 'r', encoding='utf-8') as f:
            self.documents = json.load(f)
        
        # 加载 UUID 映射
        uuid_path = load_dir / "uuid_mappings.json"
        if uuid_path.exists():
            with open(uuid_path, 'r', encoding='utf-8') as f:
                mappings = json.load(f)
                self.uuid_to_docid = mappings['uuid_to_docid']
                self.docid_to_uuid = mappings['docid_to_uuid']
        
        # 加载图数据
        with open(load_dir / "graph_data.pkl", 'rb') as f:
            data = pickle.load(f)
            self.text_chunks = data['text_chunks']
            self.chunk_to_doc = data['chunk_to_doc']
            self.entities = data['entities']
            
            # 恢复实体对齐
            if 'entity_alignments' in data:
                self.entity_alignments = {
                    k: EntityAlignment(v[0], v[1], v[2]) 
                    for k, v in data['entity_alignments'].items()
                }
            
            self.relationships = data['relationships']
            self.claims = data['claims']
            self.communities = data['communities']
            self.community_summaries = data['community_summaries']
            self.community_embeddings = data['community_embeddings']
            
            if 'needs_rebuild' in data:
                self.needs_rebuild = data['needs_rebuild']
        
        # 加载图
        with open(load_dir / "graph.gpickle", 'rb') as f:
            self.graph = pickle.load(f)
        
        # 加载 FAISS
        index_path = load_dir / "faiss_index.bin"
        if index_path.exists():
            self.community_summary_index = faiss.read_index(str(index_path))
    
    async def close(self):
        """关闭客户端连接"""
        await self.llm_client.close()
        await self.embedding_client.close()


